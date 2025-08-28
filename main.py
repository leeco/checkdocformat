from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Inches, Cm, Mm, Pt, Emu
import json
import numpy as np  # 用于统计分析

# 导入AI节点分类器
from ai_node_classifier import create_classifier

def pt_to_font_size(pt_size):
    """将磅值转换为中文标准字号"""
    font_size_map = {
        42: '初号',      # 42pt
        36: '小初',      # 36pt
        26: '一号',      # 26pt
        24: '小一',      # 24pt
        22: '二号',      # 22pt
        18: '小二',      # 18pt
        16: '三号',      # 16pt
        15: '小三',      # 15pt
        14: '四号',      # 14pt
        12: '小四',      # 12pt
        10.5: '五号',    # 10.5pt
        9: '小五',       # 9pt
        7.5: '六号',     # 7.5pt
        5.5: '小六',     # 5.5pt
        5: '七号',       # 5pt
    }
    # 找到最接近的字号
    closest_size = min(font_size_map.keys(), key=lambda x: abs(x - pt_size))
    return font_size_map[closest_size]

def format_char_value(value):
    """格式化字符数值，返回中文表示"""
    if value == 0:
        return "0字"
    elif value == int(value):
        return f"{int(value)}字"
    else:
        return f"{value}字"

def pt_to_char_accurate(pt_value, font_size=12, font_name="Default"):
    """
    根据实际字体和字号计算磅值到字符数的转换
    """
    if pt_value <= 0:
        return 0
    
    # 根据字体类型和字号计算实际字符宽度
    char_width = calculate_char_width(font_size, font_name)
    
    # 计算字符数
    char_count = pt_value / char_width
    
    # 四舍五入到0.1精度
    return round(char_count, 1)

def calculate_char_width(font_size, font_name="Default"):
    """
    根据字体和字号计算中文字符的实际宽度（磅值）
    """
    # 不同字体的字符宽度系数
    font_width_factors = {
        '宋体': 1.0,
        'SimSun': 1.0,
        '仿宋': 1.0,
        'FangSong': 1.0,
        '仿宋_GB2312': 1.0,
        '黑体': 1.0,
        'SimHei': 1.0,
        '楷体': 1.0,
        'KaiTi': 1.0,
        '微软雅黑': 0.95,
        'Microsoft YaHei': 0.95,
        'Arial': 0.6,  # 英文字体较窄
        'Times New Roman': 0.6,
        'Calibri': 0.6,
        'Default': 1.0,  # 默认按中文字体处理
    }
    
    # 获取字体宽度系数
    width_factor = font_width_factors.get(font_name, 1.0)
    
    # 在Word中，中文字符的宽度通常等于字号磅值
    # 但需要考虑字体的实际宽度特征
    base_char_width = font_size * width_factor
    
    # 对于中文字体，还需要考虑Word的内部渲染方式
    # 经验调整：Word中的字符宽度通常比理论值略小
    if font_name in ['宋体', 'SimSun', '仿宋', 'FangSong', '仿宋_GB2312', 'Default']:
        # 中文字体的经验调整系数
        base_char_width *= 0.92
    
    return base_char_width

# 全局配置：可以根据实际情况调整字符转换比例
CHAR_CONVERSION_RATIO = 11.2  # 1字符对应的磅值

def adjust_char_conversion_ratio(new_ratio):
    """调整字符转换比例"""
    global CHAR_CONVERSION_RATIO
    CHAR_CONVERSION_RATIO = new_ratio
    print(f"字符转换比例已调整为: 1字符 = {new_ratio}磅")

def get_measurement_info(measurement):
    """Get measurement value and unit from Word measurement object"""
    if measurement is None:
        return 0, '磅', 'pt'
    
    try:
        # python-docx测量对象有多种属性，优先选择最常用的
        if hasattr(measurement, 'pt') and measurement.pt is not None:
            return round(measurement.pt, 2), '磅', 'pt'
        elif hasattr(measurement, 'inches') and measurement.inches is not None:
            return round(measurement.inches, 2), '英寸', 'inch'
        elif hasattr(measurement, 'cm') and measurement.cm is not None:
            return round(measurement.cm, 2), '厘米', 'cm'
        elif hasattr(measurement, 'mm') and measurement.mm is not None:
            return round(measurement.mm, 2), '毫米', 'mm'
        elif hasattr(measurement, 'emu') and measurement.emu is not None:
            # EMU转换为磅 (1 point = 12700 EMU)
            pt_value = measurement.emu / 12700
            return round(pt_value, 2), '磅', 'pt'
        else:
            # 如果是数字，假设是磅
            try:
                return round(float(measurement), 2), '磅', 'pt'
            except (TypeError, ValueError):
                return 0, '磅', 'pt'
    except Exception:
        return 0, '磅', 'pt'

def get_alignment_value(alignment):
    """获取对齐方式的中文描述"""
    if alignment is None:
        return '左对齐', 'None'
    
    try:
        # 首先尝试直接比较枚举值
        if alignment == WD_ALIGN_PARAGRAPH.LEFT:
            return '左对齐', 'LEFT (0)'
        elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return '居中', 'CENTER (1)'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return '右对齐', 'RIGHT (2)'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return '两端对齐', 'JUSTIFY (3)'
        elif alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
            return '分散对齐', 'DISTRIBUTE (4)'
        elif alignment == WD_ALIGN_PARAGRAPH.THAI_JUSTIFY:
            return '泰文两端对齐', 'THAI_JUSTIFY (5)'
        elif hasattr(WD_ALIGN_PARAGRAPH, 'JUSTIFY_MED') and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY_MED:
            return '两端对齐(中等)', 'JUSTIFY_MED (6)'
        elif hasattr(WD_ALIGN_PARAGRAPH, 'JUSTIFY_HI') and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY_HI:
            return '两端对齐(高)', 'JUSTIFY_HI (7)'
        elif hasattr(WD_ALIGN_PARAGRAPH, 'JUSTIFY_LOW') and alignment == WD_ALIGN_PARAGRAPH.JUSTIFY_LOW:
            return '两端对齐(低)', 'JUSTIFY_LOW (8)'
    except Exception:
        pass
    
    # 如果枚举值比较失败，尝试数字值比较
    try:
        alignment_int = int(alignment)
        alignment_map = {
            0: ('左对齐', 'LEFT (0)'),
            1: ('居中', 'CENTER (1)'),
            2: ('右对齐', 'RIGHT (2)'),
            3: ('两端对齐', 'JUSTIFY (3)'),
            4: ('分散对齐', 'DISTRIBUTE (4)'),
            5: ('泰文两端对齐', 'THAI_JUSTIFY (5)'),
            6: ('两端对齐(中等)', 'JUSTIFY_MED (6)'),
            7: ('两端对齐(高)', 'JUSTIFY_HI (7)'),
            8: ('两端对齐(低)', 'JUSTIFY_LOW (8)')
        }
        
        if alignment_int in alignment_map:
            return alignment_map[alignment_int]
    except (ValueError, TypeError):
        pass
    
    # 最后尝试字符串匹配
    alignment_str = str(alignment)
    string_map = {
        'CENTER (1)': ('居中', 'CENTER (1)'),
        'LEFT (0)': ('左对齐', 'LEFT (0)'), 
        'RIGHT (2)': ('右对齐', 'RIGHT (2)'),
        'JUSTIFY (3)': ('两端对齐', 'JUSTIFY (3)'),
        'DISTRIBUTE (4)': ('分散对齐', 'DISTRIBUTE (4)'),
        'THAI_JUSTIFY (5)': ('泰文两端对齐', 'THAI_JUSTIFY (5)'),
        'JUSTIFY_MED (6)': ('两端对齐(中等)', 'JUSTIFY_MED (6)'),
        'JUSTIFY_HI (7)': ('两端对齐(高)', 'JUSTIFY_HI (7)'),
        'JUSTIFY_LOW (8)': ('两端对齐(低)', 'JUSTIFY_LOW (8)')
    }
    
    if alignment_str in string_map:
        return string_map[alignment_str]
    
    # 如果都不匹配，返回原始值
    return f'未知对齐方式({alignment_str})', str(alignment)

class Node:
    def __init__(self, node_type, content, font=None, size=None, bold=False, paragraph_format=None):
        self.type = node_type  # e.g., 'document_title', 'heading1', 'paragraph', 'list_item', etc.
        self.content = content
        # 字体信息
        self.font = font
        self.size = size  # in pt
        self.font_size_name = pt_to_font_size(size) if size else None  # 字号名称
        self.bold = bold
        # 完整格式信息（包含所有段落格式）
        self.paragraph_format = paragraph_format or {}
        self.children = []  # list of child nodes

    def add_child(self, child_node):
        self.children.append(child_node)
    
    # 属性访问器 - 通过paragraph_format获取格式信息
    @property
    def spacing(self):
        """段后间距"""
        return self.paragraph_format.get('space_after', {}).get('value', 0)
    
    @property 
    def line_spacing(self):
        """行距"""
        line_spacing_info = self.paragraph_format.get('line_spacing', {})
        return line_spacing_info.get('value') if isinstance(line_spacing_info, dict) else None
    
    @property
    def indentation(self):
        """缩进信息"""
        indent_info = {}
        if 'left_indent' in self.paragraph_format:
            indent_info['left'] = self.paragraph_format['left_indent']
        if 'right_indent' in self.paragraph_format:
            indent_info['right'] = self.paragraph_format['right_indent']
        if 'first_line_indent' in self.paragraph_format:
            indent_info['first_line'] = self.paragraph_format['first_line_indent']
        if 'hanging_indent' in self.paragraph_format:
            indent_info['hanging'] = self.paragraph_format['hanging_indent']
        return indent_info
    
    @property
    def outline_level(self):
        """大纲级别"""
        return self.paragraph_format.get('outline_level')
    
    @property
    def alignment(self):
        """对齐方式"""
        return self.paragraph_format.get('alignment')
    
    @property
    def direction(self):
        """文字方向"""
        return self.paragraph_format.get('direction')

    def to_dict(self):
        return {
            'type': self.type,
            'content': self.content,
            'font': self.font,
            'size': self.size,
            'font_size_name': self.font_size_name,
            'bold': self.bold,
            # 保持向后兼容，同时提供属性访问
            'spacing': self.spacing,
            'line_spacing': self.line_spacing,
            'indentation': self.indentation,
            'outline_level': self.outline_level,
            'alignment': self.alignment,
            'direction': self.direction,
            'paragraph_format': self.paragraph_format,
            'children': [child.to_dict() for child in self.children]
        }

def get_font_size_distribution(doc):
    """Collect font sizes from all paragraphs."""
    sizes = []
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():  # Skip empty paragraphs
            continue
        first_run = paragraph.runs[0] if paragraph.runs else None
        size = first_run.font.size.pt if first_run and first_run.font.size else 12.0
        sizes.append(size)
    return sizes

def extract_paragraph_formatting(paragraph, font_name="Default", font_size=12):
    """Extract detailed formatting information from paragraph including all properties from the dialog."""
    format_info = {}
    
    # Get paragraph format
    pf = paragraph.paragraph_format
    
    # 1. 常规 (General) Section
    # Alignment (对齐方式) - 修复后的逻辑
    alignment_value, alignment_raw = get_alignment_value(pf.alignment)
    format_info['alignment'] = {
            'type': '对齐方式',
        'value': alignment_value,
        'raw_value': alignment_raw
        }
    
    # Outline Level (大纲级别)
    outline_map = {
        0: '标题1',
        1: '标题2', 
        2: '标题3',
        3: '标题4',
        4: '标题5',
        5: '标题6',
        6: '标题7',
        7: '标题8',
        8: '标题9',
        9: '正文文本'
    }
    
    if hasattr(pf, 'outline_level') and pf.outline_level is not None:
        format_info['outline_level'] = {
            'type': '大纲级别',
            'value': outline_map.get(pf.outline_level, str(pf.outline_level)),
            'raw_value': pf.outline_level
        }
    else:
        format_info['outline_level'] = {
            'type': '大纲级别',
            'value': '正文文本',
            'raw_value': 9
        }
    
    # Direction (方向) - 从右向左/从左向右
    format_info['direction'] = {
        'type': '方向',
        'value': '从左向右',
        'raw_value': 'LTR'
    }
    
    # 2. 缩进 (Indentation) Section
    def get_indent_info(indent_value, indent_type, font_size=12, font_name="Default"):
        """获取缩进信息，根据实际字体和字号转换为中文字符单位"""
        if indent_value is None:
            return {
                'type': indent_type,
            'value': 0,
                'unit': '字',
                'unit_code': 'char',
                'display_text': '0字',
                'pt_value': 0,
                'font_size': font_size,
                'font_name': font_name,
                'conversion_ratio': '无缩进',
                'raw_value': None
            }
        
        # 获取磅值
        pt_value, _, _ = get_measurement_info(indent_value)
        
        # 根据实际字体和字号进行字符转换
        char_value = pt_to_char_accurate(pt_value, font_size, font_name)
        
        # 格式化显示，如果是整数就不显示小数点
        if char_value == int(char_value):
            display_value = int(char_value)
        else:
            display_value = char_value
        
        # 计算实际的字符宽度用于显示
        actual_char_width = calculate_char_width(font_size, font_name)
        
        return {
            'type': indent_type,
            'value': display_value,
            'unit': '字',
            'unit_code': 'char',
            'display_text': format_char_value(display_value),  # 中文显示格式
            'pt_value': pt_value,  # 保留原始磅值用于参考
            'font_size': font_size,
            'font_name': font_name,
            'char_width': round(actual_char_width, 1),  # 基于字体字号的字符宽度
            'conversion_ratio': f'1字={actual_char_width:.1f}磅({font_name},{font_size}号)',
            'raw_value': indent_value
        }
    
    # 左缩进
    format_info['left_indent'] = get_indent_info(pf.left_indent, '左缩进', font_size, font_name)
    
    # 右缩进
    format_info['right_indent'] = get_indent_info(pf.right_indent, '右缩进', font_size, font_name)
    
    # 首行缩进/悬挂缩进
    if pf.first_line_indent:
        pt_value, _, _ = get_measurement_info(pf.first_line_indent)
        
        # 根据实际字体和字号进行字符转换
        char_value = pt_to_char_accurate(pt_value, font_size, font_name)
        actual_char_width = calculate_char_width(font_size, font_name)
        
        # 格式化显示值
        if char_value == int(char_value):
            display_value = int(char_value)
        else:
            display_value = char_value
        
        if pt_value > 0:
            format_info['first_line_indent'] = {
                'type': '首行缩进',
                'value': display_value,
                'unit': '字',
                'unit_code': 'char',
                'display_text': format_char_value(display_value),
                'pt_value': pt_value,
                'font_size': font_size,
                'font_name': font_name,
                'char_width': round(actual_char_width, 1),
                'conversion_ratio': f'1字={actual_char_width:.1f}磅({font_name},{font_size}号)',
                'raw_value': pf.first_line_indent
            }
        elif pt_value < 0:
            format_info['hanging_indent'] = {
                'type': '悬挂缩进',
                'value': abs(display_value),
                'unit': '字',
                'unit_code': 'char',
                'display_text': format_char_value(abs(display_value)),
                'pt_value': abs(pt_value),
                'font_size': font_size,
                'font_name': font_name,
                'char_width': round(actual_char_width, 1),
                'conversion_ratio': f'1字={actual_char_width:.1f}磅({font_name},{font_size}号)',
                'raw_value': pf.first_line_indent
            }
        else:
            format_info['first_line_indent'] = {
                'type': '首行缩进',
                'value': 0,
                'unit': '字',
                'unit_code': 'char',
                'display_text': '0字',
                'pt_value': 0,
                'font_size': font_size,
                'font_name': font_name,
                'char_width': round(actual_char_width, 1),
                'conversion_ratio': f'1字={actual_char_width:.1f}磅({font_name},{font_size}号)',
                'raw_value': pf.first_line_indent
            }
    else:
        default_char_width = calculate_char_width(font_size, font_name)
        format_info['first_line_indent'] = {
            'type': '首行缩进',
            'value': 0,
            'unit': '字',
            'unit_code': 'char',
            'display_text': '0字',
            'pt_value': 0,
            'font_size': font_size,
            'font_name': font_name,
            'char_width': round(default_char_width, 1),
            'conversion_ratio': f'1字={default_char_width:.1f}磅({font_name},{font_size}号)',
            'raw_value': None
        }
    
    # 3. 间距 (Spacing) Section
    # 段前 (Before paragraph)
    if pf.space_before:
        value, unit_name, unit_code = get_measurement_info(pf.space_before)
        format_info['space_before'] = {
            'type': '段前间距',
            'value': value,
            'unit': unit_name,
            'unit_code': unit_code
        }
    else:
        format_info['space_before'] = {
            'type': '段前间距',
            'value': 0,
            'unit': '磅',
            'unit_code': 'pt'
        }
    
    # 段后 (After paragraph)
    if pf.space_after:
        value, unit_name, unit_code = get_measurement_info(pf.space_after)
        format_info['space_after'] = {
            'type': '段后间距',
            'value': value,
            'unit': unit_name,
            'unit_code': unit_code
        }
    else:
        format_info['space_after'] = {
            'type': '段后间距',
            'value': 0,
            'unit': '磅',
            'unit_code': 'pt'
        }
    
    # 行距 (Line spacing) - 处理不同类型的行距
    def get_line_spacing_info(paragraph_format):
        """获取行距信息"""
        try:
            line_spacing = paragraph_format.line_spacing
            line_spacing_rule = paragraph_format.line_spacing_rule
            
            # 行距规则映射
            rule_map = {
                0: ('单倍行距', 'SINGLE'),      # WD_LINE_SPACING.SINGLE
                1: ('1.5倍行距', 'ONE_POINT_FIVE'),  # WD_LINE_SPACING.ONE_POINT_FIVE  
                2: ('2倍行距', 'DOUBLE'),        # WD_LINE_SPACING.DOUBLE
                3: ('最小值', 'AT_LEAST'),       # WD_LINE_SPACING.AT_LEAST
                4: ('固定值', 'EXACTLY'),        # WD_LINE_SPACING.EXACTLY
                5: ('多倍行距', 'MULTIPLE')      # WD_LINE_SPACING.MULTIPLE
            }
            
            if line_spacing is None and line_spacing_rule is None:
                return {
                    'type': '行距',
                    'rule': '单倍行距',
                    'value': 1.0,
                    'unit': '倍数',
                    'unit_code': 'multiple',
                    'raw_rule': None,
                    'raw_value': None
                }
            
            # 获取规则名称
            if line_spacing_rule is not None:
                rule_int = int(line_spacing_rule) if isinstance(line_spacing_rule, int) else line_spacing_rule
                rule_name, rule_code = rule_map.get(rule_int, ('未知规则', str(line_spacing_rule)))
            else:
                rule_name, rule_code = '自动', 'AUTO'
            
            # 处理行距值
            if line_spacing is not None:
                # 对于倍数行距（SINGLE, ONE_POINT_FIVE, DOUBLE, MULTIPLE）
                if line_spacing_rule in [0, 1, 2, 5]:  # SINGLE, ONE_POINT_FIVE, DOUBLE, MULTIPLE
                    if line_spacing_rule == 0:  # SINGLE
                        value = 1.0
                    elif line_spacing_rule == 1:  # ONE_POINT_FIVE
                        value = 1.5
                    elif line_spacing_rule == 2:  # DOUBLE
                        value = 2.0
                    else:  # MULTIPLE
                        # 对于多倍行距，line_spacing就是倍数
                        value = round(float(line_spacing), 2)
                    
                    return {
                        'type': '行距',
                        'rule': rule_name,
                        'value': value,
                        'unit': '倍数',
                        'unit_code': 'multiple',
                        'raw_rule': line_spacing_rule,
                        'raw_value': line_spacing
                    }
                
                # 对于固定值和最小值（EXACTLY, AT_LEAST）
                elif line_spacing_rule in [3, 4]:  # AT_LEAST, EXACTLY
                    # line_spacing是Length对象，需要转换
                    value, unit_name, unit_code = get_measurement_info(line_spacing)
                    return {
                        'type': '行距',
                        'rule': rule_name,
                        'value': value,
                        'unit': unit_name,
                        'unit_code': unit_code,
                        'raw_rule': line_spacing_rule,
                        'raw_value': line_spacing
                    }
                else:
                    # 其他情况，尝试作为倍数处理
                    try:
                        value = round(float(line_spacing), 2)
                        return {
                            'type': '行距',
                            'rule': rule_name,
                            'value': value,
                            'unit': '倍数',
                            'unit_code': 'multiple',
                            'raw_rule': line_spacing_rule,
                            'raw_value': line_spacing
                        }
                    except (TypeError, ValueError):
                        # 如果无法转换为数字，尝试作为Length处理
                        value, unit_name, unit_code = get_measurement_info(line_spacing)
                        return {
                            'type': '行距',
                            'rule': rule_name,
                            'value': value,
                            'unit': unit_name,
                            'unit_code': unit_code,
                            'raw_rule': line_spacing_rule,
                            'raw_value': line_spacing
                }
            else:
                # 只有规则没有值的情况
                default_values = {
                    0: 1.0,   # SINGLE
                    1: 1.5,   # ONE_POINT_FIVE
                    2: 2.0,   # DOUBLE
                }
                value = default_values.get(line_spacing_rule, 1.0)
                return {
                    'type': '行距',
                    'rule': rule_name,
                    'value': value,
                    'unit': '倍数',
                    'unit_code': 'multiple',
                    'raw_rule': line_spacing_rule,
                    'raw_value': line_spacing
                }
                
        except Exception as e:
            # 异常处理，返回默认值
            return {
            'type': '行距',
            'rule': '单倍行距',
            'value': 1.0,
            'unit': '倍数',
                'unit_code': 'multiple',
                'raw_rule': None,
                'raw_value': None,
                'error': str(e)
        }
    
    format_info['line_spacing'] = get_line_spacing_info(pf)
    
    return format_info

def parse_docx_to_tree(file_path, api_key: str = None):
    """解析Word文档为树形结构，使用AI分类器"""
    doc = Document(file_path)
    root = Node('root', 'Document Root')
    stack = [root]
    
    # 创建AI节点分类器
    classifier = create_classifier(api_key)
    
    # 收集所有段落用于上下文分析（包括空行）
    all_paragraphs = []
    for paragraph in doc.paragraphs:
        all_paragraphs.append(paragraph)
    
    # Step 2: Parse paragraphs and build tree
    for i, paragraph in enumerate(all_paragraphs):
        
        # 检查是否为空行
        content = paragraph.text.strip()
        is_empty_line = not content
        
        if is_empty_line:
            # 为空行创建特殊节点，使用默认字体信息
            format_info = extract_paragraph_formatting(paragraph, "Default", 12.0)
            
            # 创建空行节点
            node_type = 'empty_line'
            font = 'Default'
            size = 12.0
            bold = False
            
            new_node = Node(
                node_type, '[空行]', font, size, bold, format_info
            )
        else:
            # 处理非空行
            # Extract properties
            first_run = paragraph.runs[0] if paragraph.runs else None
            font = first_run.font.name if first_run and first_run.font.name else 'Default'
            size = first_run.font.size.pt if first_run and first_run.font.size else 12.0
            bold = first_run.font.bold if first_run and first_run.font.bold else False
            
            # Get paragraph formatting，传入实际的字体和字号信息
            format_info = extract_paragraph_formatting(paragraph, font, size)
            
            # 获取上下文节点（前3个段落，只包含非空行）
            context_nodes = []
            for j in range(max(0, i-3), i):
                if j < len(all_paragraphs) and all_paragraphs[j].text.strip():
                    context_nodes.append({
                        'content': all_paragraphs[j].text.strip(),
                        'type': 'paragraph'  # 临时类型
                    })
            
            # 使用AI分类器
            node_type = classifier.classify_node(
                content=content,
                font=font,
                size=size,
                bold=bold,
                outline_level=format_info.get('outline_level'),
                alignment=format_info.get('alignment'),
                paragraph_format=format_info,
                context_nodes=context_nodes
            )
            
            new_node = Node(
                node_type, content, font, size, bold, format_info
            )
        
        # Build tree: Adjust stack based on level
        # 更新层级映射，包含所有支持的节点类型
        level_map = {
            '发文标题': 0,        # 发文标题层级最高
            '主送机关': 1,        # 主送机关
            '一级标题': 2,        # 一级标题
            '二级标题': 3,        # 二级标题  
            '三级标题': 4,        # 三级标题
            '四级标题': 5,        # 四级标题
            '列表项': 6,          # 列表项
            '普通段落': 7,        # 普通段落
            '结尾': 8,            # 结尾
            '落款': 9,            # 落款
            '附件': 10,           # 附件
            '分隔符': 11,         # 分隔符
            '空行': 12            # 空行层级最低，不影响文档结构
        }
        
        current_level = level_map.get(node_type, 6)
        while len(stack) > 1 and current_level <= level_map.get(stack[-1].type, 6):
            stack.pop()
        stack[-1].add_child(new_node)
        stack.append(new_node)
    
    return root

# Example usage
if __name__ == "__main__":
    file_path = '1.docx'  # Replace with your file path
    
    # 解析文档
    tree_root = parse_docx_to_tree(file_path)

    # Print tree as JSON
    print(json.dumps(tree_root.to_dict(), indent=4, ensure_ascii=False))

    # Print as tree string
    def print_tree(node, indent=0):
        spacing_info = f"spacing={node.spacing}" if node.spacing else ""
        line_spacing_info = f"line_spacing={node.line_spacing}" if node.line_spacing else ""
        indent_info = f"indent={node.indentation}" if node.indentation else ""
        alignment_info = f"alignment={node.alignment}" if node.alignment else ""
        outline_info = f"outline={node.outline_level}" if node.outline_level else ""
        font_size_info = f"字号={node.font_size_name}" if node.font_size_name else ""
        
        info_parts = [f"font={node.font}", f"size={node.size}pt", f"bold={node.bold}"]
        if font_size_info:
            info_parts.append(font_size_info)
        if spacing_info:
            info_parts.append(spacing_info)
        if line_spacing_info:
            info_parts.append(line_spacing_info)
        if indent_info:
            info_parts.append(indent_info)
        if alignment_info:
            info_parts.append(alignment_info)
        if outline_info:
            info_parts.append(outline_info)
        
        info_str = ", ".join(info_parts)
        print('  ' * indent + f"{node.type}: {node.content} ({info_str})")
        for child in node.children:
            print_tree(child, indent + 1)

    print("\n" + "="*50)
    print("TREE STRUCTURE:")
    print("="*50)
    print_tree(tree_root)

    # Save to JSON file
    with open('tree_output.json', 'w', encoding='utf-8') as f:
        json.dump(tree_root.to_dict(), f, ensure_ascii=False, indent=4)